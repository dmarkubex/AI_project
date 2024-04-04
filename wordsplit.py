import os
import re
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx.oxml.ns import qn
from docx.oxml import parse_xml

os.environ['base_url'] = "https://apivip.aiproxy.io"
os.environ['OPENAI_API_KEY'] = "sk-TKeAI4nyygMazxskGgxWaeibhzxFPyzGeCVTzsg4xoPEZcoc"

def remove_header_footer(text):
    """
    Removes headers and footers from text based on certain patterns.

    Args:
    text (str): Text with headers and footers.

    Returns:
    str: Cleaned text without headers and footers.
    """
    lines = text.split("\n")
    cleaned_lines = [line for line in lines if not re.match(r'^\s*-\s*\d+\s*-\s*$', line)]
    return "\n".join(cleaned_lines)

def extract_text_from_docx(doc):
    """
    Extracts text from a Word document, including paragraphs.

    Args:
    doc (Document): Word document object.

    Returns:
    str: Extracted text from the document.
    """
    paragraphs_text = '\n'.join([para.text for para in doc.paragraphs])
    return paragraphs_text

def extract_tables_from_docx(doc, output_file_path):
    """
    Extracts tables from a Word document and saves them to another Word document.
    """
    tables_doc = Document()
    for table in doc.tables:
        # 直接复制表格的XML
        tbl_xml = table._element.xml
        new_tbl = parse_xml(tbl_xml)
        tables_doc.element.body.append(new_tbl)

    tables_doc.save(output_file_path)

def process_single_docx(docx_path, output_dir):
    """
    Processes a single Word document, handling recursive character-based text splitting and table extraction.
    """
    doc = Document(docx_path)
    extracted_text = extract_text_from_docx(doc)
    cleaned_text = remove_header_footer(extracted_text)

    # 使用 RecursiveCharacterTextSplitter 进行文本分割，并处理中文字符
    text_splitter = RecursiveCharacterTextSplitter(
        separators=[
            "\n\n",
            "\n",
            " ",
            ".",
            ",",
            "\u200B",  # Zero-width space
            "\uff0c",  # Fullwidth comma
            "\u3001",  # Ideographic comma
            "\uff0e",  # Fullwidth full stop
            "\u3002",  # Ideographic full stop
            "",
        ],
        chunk_size=300,
        chunk_overlap=60,
        length_function=len,
        is_separator_regex=False
    )
    split_texts = text_splitter.split_text(cleaned_text)

    processed_text = '&&&&&&'.join(split_texts)

    # 保存处理后的文本到文本文件
    text_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".txt"
    with open(os.path.join(output_dir, text_filename), 'w', encoding='utf-8') as f:
        f.write(processed_text)

    # 提取并保存原始文档中的表格到新的Word文档
    tables_filename = os.path.splitext(os.path.basename(docx_path))[0] + "-table.docx"
    extract_tables_from_docx(doc, os.path.join(output_dir, tables_filename))



def process_all_docx(input_dir, output_dir):
    """
    Processes all Word documents in the specified directory.
    """
    for filename in os.listdir(input_dir):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            docx_path = os.path.join(input_dir, filename)
            process_single_docx(docx_path, output_dir)

def create_directory_if_not_exists(directory):
    """
    Creates a directory if it does not already exist.

    Args:
    directory (str): The directory to create.
    """
    if not os.path.exists(directory):
        os.makedirs(directory)

output_dir = "C:/Users/刁敏/Documents/Project/AI/output"
create_directory_if_not_exists(output_dir)

# 示例使用
input_dir = "C:/Users/刁敏/Documents/Project/AI/word"
output_dir = "C:/Users/刁敏/Documents/Project/AI/output"
process_all_docx(input_dir, output_dir)